from PyQt5.QtWidgets import QMenu, QAction, QFileDialog, QMessageBox
from PyQt5.QtCore import QTimer
from docx import Document
from docx.table import Table
import getDocContent
import logging


class FileMenu(QMenu):
    """Class that provides standard functions to save and to open files"""
    def __init__(self, app, mainWindow, parent, centralW):
        QMenu.__init__(self, parent)
        logging.basicConfig(level=logging.INFO)
        self.logger = logging.getLogger(__name__)

        self.app = app
        self.doc = centralW.getDocument()
        self.cursor = centralW.getCursor()
        self.parent = parent
        self.mainWindow = mainWindow

        self.setTitle('Файл')
        parent.addMenu(self)

        self.__newAction()
        self.__openAction()
        self.__saveAction()
        self.__saveAsAction()
        self.__exitAction()

        self.fileName = None
        self.star = False

        self.timer = QTimer()
        self.timer.timeout.connect(self.__changeTitleIfModified)
        self.timer.start(50)

    def __newAction(self):
        self.newAction = QAction('Новый', self)
        self.newAction.setShortcut('Ctrl+N')
        self.newAction.triggered.connect(self.__new)
        self.addAction(self.newAction)
        self.addSeparator()

    def __openAction(self):
        self.openAction = QAction('Открыть', self)
        self.openAction.setShortcut('Ctrl+O')
        self.openAction.triggered.connect(self.__open)
        self.addAction(self.openAction)

    def __saveAction(self):
        self.saveAction = QAction('Сохранить', self)
        self.saveAction.setShortcut('Ctrl+S')
        self.saveAction.triggered.connect(self.__save)
        self.addAction(self.saveAction)

    def __saveAsAction(self):
        self.saveAsAction = QAction('Сохранить как...', self)
        self.saveAsAction.setShortcut('Ctrl+Shift+S')
        self.saveAsAction.triggered.connect(self.__saveAs)
        self.addAction(self.saveAsAction)
        self.addSeparator()

    def __exitAction(self):
        self.exitAction = QAction('Выйти', self)
        self.exitAction.setShortcut('Ctrl+Q')
        self.exitAction.triggered.connect(self.__exit)
        self.addAction(self.exitAction)

    def __open(self):
        self.saveIfChanged(self.__openFunction)

    def __openFunction(self):
        self.fileName = QFileDialog.getOpenFileName(self.parent, '', '', "Document Microsoft Word (*.docx)")

        if self.fileName is not None and self.fileName[0]:
            doc = Document(self.fileName[0])

            obj = getDocContent.iter_block_items(doc)

            for i in obj:
                if type(i) == Table:
                    self.cursor.insertText('\n')
                    for row_index in range(len(i.rows)):
                        for column_index in range(len(i.columns)):
                            self.cursor.insertText(i.cell(row_index, column_index).text + '\t')
                        self.cursor.insertText('\n')
                    self.cursor.insertText('\n')
                else:
                    self.cursor.insertText(i.text + '\n')

        self.doc.setModified(False)
        self.mainWindow.setWindowTitle(self.mainWindow.windowTitle() + '-' + self.fileName[0])

    def __saveAs(self):
        self.fileName = QFileDialog.getSaveFileName(self.parent, '', '', "Document Microsoft Word (*.docx)")

        if self.fileName is not None and self.fileName[0]:
            self.__addToDocx()

    def __new(self):
        self.saveIfChanged()

    def __exit(self):
        self.saveIfChanged(self.__exitFunction)

    def __exitFunction(self):
        self.app.quit()

    def __message(self, windTitle, text, infoText):
        info = QMessageBox()
        info.setIcon(QMessageBox.Information)
        info.setWindowTitle(windTitle)
        info.setText(text)
        info.setInformativeText(infoText)
        info.setStandardButtons(QMessageBox.Yes | QMessageBox.No | QMessageBox.Cancel)

        return hex(info.exec_())

    def saveIfChanged(self, function=None):
        if self.doc.isModified():
            retval = self.__message('Изменение текста', 'Файл изменен', 'Вы хотите сохранить изменения')
            if retval == hex(0x10000):
                self.doc.clear()
                self.doc.setModified(False)
                self.mainWindow.setWindowTitle('Корректор')
                self.fileName = None
                self.star = False
                if function is not None:
                    function()
                return True
            elif retval == hex(0x4000):
                self.__save()
                if function is not None:
                    function()
                return True
            elif retval == hex(0x400000):
                return False
        else:
            self.doc.clear()
            self.doc.setModified(False)
            self.mainWindow.setWindowTitle('Корректор')
            self.fileName = None
            if function is not None:
                function()
            return True

    def __save(self):
        if self.fileName is not None and self.fileName[0]:
            self.__addToDocx()
        else:
            self.__saveAs()

    def __addToDocx(self):
        doc = Document()
        for val in self.doc.toPlainText().split('\n'):
            doc.add_paragraph(val + '\n')

        fileName = self.fileName[0].strip('.docx')

        doc.save(fileName + '.docx')
        self.doc.setModified(False)
        self.star = False
        self.mainWindow.setWindowTitle('Корректор' + '-' + self.fileName[0])

    def __changeTitleIfModified(self):
        if self.doc.isModified() and not self.star:
            self.mainWindow.setWindowTitle(self.mainWindow.windowTitle() + '*')
            self.star = True
